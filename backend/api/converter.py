from pdf2image import convert_from_path
import pytesseract
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from pdf2image import convert_from_path
from docx import Document
from docx.enum.text import WD_BREAK

class Converter:

    @staticmethod
    def wrap_text(text, max_length):
        """Wrap text to limit the line length and return a list of strings."""
        words = text.split()
        lines = []
        current_line = ""

        for word in words:
            if len(current_line) + len(word) + 1 <= max_length:
                current_line += word + " "
            else:
                lines.append(current_line.strip())
                current_line = word + " "

        if current_line:
            lines.append(current_line.strip())

        return lines

    @staticmethod
    def pdf_to_text(pdf_path):
        """
        Extract text content from a PDF file.

        Parameters:
        - pdf_path (str): The path to the PDF file.

        Returns:
        str: The extracted text content from the PDF.

        Note:
        - This method uses Poppler for PDF rendering and Tesseract OCR for text extraction.
        - Ensure that the 'poppler_path' and 'tesseract_cmd' paths are correctly set for your environment.
        - The method returns the extracted text content from the PDF.
        """

        poppler_path = r'poppler-23.11.0\Library\bin'
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        images = convert_from_path(pdf_path, poppler_path=poppler_path)
        text = ''
        for i, img in enumerate(images):
            text += f'\nPage {i + 1}:\n\n'
            text += pytesseract.image_to_string(img, lang='eng')

        return text
    
    # If ever custom file naming convention is necessary
    # added username assessment id, and assessment_name for the file name to avoid errors
    # file naming convention: {assessment_name}_quiz_{username}-{assessment_id}.pdf
    # def quiz_to_pdf(assessment, username, assessment_id, assessment_name, type):
    @staticmethod
    def quiz_to_pdf(assessment, type, name, username):
        """
        Convert a quiz assessment in JSON format to a PDF document.

        Parameters:
        - assessment (dict): The quiz assessment in dictionary format.
        - type (str): The type of assessment ("multiple choice", "identification", "true or false", "fill in the blanks", "essay").

        Note:
        - This method creates a PDF document with formatted content based on the quiz assessment.
        - The PDF is saved to the 'Project Files' directory with the naming convention '{name}_quiz.pdf'.
        - The method supports different types of assessments, each with specific formatting.
        - The assessment dictionary should have the structure consistent with the expected format for the given type.
        """

        print("assessment: ", assessment)

        # Create a PDF document
        pdf_canvas = canvas.Canvas(rf"backend\api\media\{username}\exports\{name}_quiz.pdf", pagesize=letter)

        pdf_canvas.setFont("Helvetica-Bold", 14)
        pdf_canvas.drawString(50, 770, f"{type}")
        pdf_canvas.setFont("Helvetica", 12)
        
        # Extract information from the JSON
        questions = assessment.get("questions", [])

        # Add content to the PDF
        pdf_canvas.setFont("Helvetica", 12)

        y_position = 750  # Adjusted starting position
        x_position = 50  # Adjusted starting position on the x-axis
        line_height = 10  # Adjusted line height
        max_line_length = 80  # Adjusted maximum line length

        for index, question in enumerate(questions, start=1):
            y_position -= 2 * line_height  # Adjust the vertical position for each question

            question_text = question.get("question", "")

            if type == "multiple choice":
                options = question.get("options", [])
                pdf_canvas.drawString(x_position, y_position, f"{index}. {question_text}")

                # Add options to the PDF
                for option_index, option in enumerate(options, start=1):
                    y_position -= line_height  # Adjust the vertical position for each option
                    wrapped_option_lines = Converter.wrap_text(f"{chr(96 + option_index)}. {option}", max_line_length)
                    for i, wrapped_option_line in enumerate(wrapped_option_lines):
                        if i == 0:
                            y_position -= line_height
                            pdf_canvas.drawString(x_position + 15, y_position, wrapped_option_line)
                        else:
                            pdf_canvas.drawString(x_position + 28, y_position - 2, wrapped_option_line)
                        y_position -= line_height

            elif type == "identification":
                wrapped_text_lines = Converter.wrap_text(f"{index}. {question_text}", max_line_length)
                for i, wrapped_text_line in enumerate(wrapped_text_lines):
                    if i == 0:
                        pdf_canvas.drawString(x_position, y_position, wrapped_text_line)
                    else:
                        pdf_canvas.drawString(x_position + 15, y_position, wrapped_text_line)
                    y_position -= line_height

            elif type == "true or false":
                wrapped_text_lines = Converter.wrap_text(f"{index}. {question_text}", max_line_length)
                for i, wrapped_text_line in enumerate(wrapped_text_lines):
                    if i == 0:
                        pdf_canvas.drawString(x_position, y_position, wrapped_text_line)
                    else:
                        pdf_canvas.drawString(x_position + 15, y_position, wrapped_text_line)
                    y_position -= line_height

            elif type == "fill in the blanks":
                lines = Converter.wrap_text(f"{index}. {question_text}", max_line_length)
                for i, line in enumerate(lines):
                    if i == 0:
                        pdf_canvas.drawString(x_position, y_position, line)
                    else:
                        pdf_canvas.drawString(x_position + 15, y_position-5, line)
                    y_position -= line_height

            elif type == "essay":
                lines = Converter.wrap_text(f"{index}. {question_text}", max_line_length)
                for i, line in enumerate(lines):
                    if i == 0:
                        pdf_canvas.drawString(x_position, y_position, line)
                    else:
                        pdf_canvas.drawString(x_position + 15, y_position-5, line)
                    y_position -= line_height

            y_position -= line_height  # Adjust the vertical position for the next question

        # Save the PDF
        pdf_canvas.save()
    
    # If ever custom file naming convention is necessary
    # added username assessment id, and assessment_name for the file name to avoid errors
    # file naming convention: {assessment_name}_quiz-answer-key_{username}-{assessment_id}.pdf
    # def quiz_answer_key(assessment, username, assessment_id, assessment_name, type):
    @staticmethod
    def quiz_answer_key(assessment, type, name, username):
        """
        Generate an answer key PDF for a quiz assessment.

        Parameters:
        - assessment (dict): The quiz assessment in dictionary format.
        - type (str): The type of assessment ("multiple choice", "identification", "true or false", "fill in the blanks", "essay").

        Note:
        - This method creates an answer key PDF document with correct answers for the quiz assessment.
        - The PDF is saved to the 'Files' directory with the naming convention '{name}_quiz_answer-key.pdf'.
        - The method supports different types of assessments, each with specific formatting for correct answers.
        - The assessment dictionary should have the structure consistent with the expected format for the given type.
        """

        # Get the questions from the assessment
        questions = assessment.get("questions", [])

        # Create a PDF document for the answer key
        pdf_canvas = canvas.Canvas(rf"backend\api\media\{username}\exports\{name}_quiz_answer-key.pdf", pagesize=letter)

        # Add content to the PDF
        pdf_canvas.setFont("Helvetica", 12)

        y_position = 780
        for index, question in enumerate(questions, start=1):
            y_position -= 15  # Adjust the vertical position for each question
            
            correct_answer = f"Question {index}: {question['answer']}"

            # this seems redundant
            # if type == "multiple choice":
            #     correct_answer = f"Question {index}: {question['answer']}"
            # elif type == "identification":
            #     correct_answer = f"Question {index}: {question['answer']}"
            # elif type == "true or false":
            #     correct_answer = f"Question {index}: {question['answer']}"
            # elif type == "fill in the blanks":
            #     correct_answer = f"Question {index}: {question['answer']}"

            print(correct_answer)
            pdf_canvas.drawString(120, y_position, correct_answer)

            y_position -= 25  # Adjust the vertical position for the next question

        # Save the PDF
        pdf_canvas.save()

    # If ever custom file naming convention is necessary
    # added username assessment id, and assessment_name for the file name to avoid errors
    # file naming convention: {assessment_name}_exam_{username}-{assessment_id}.pdf
    # def exam_to_pdf(exam, username, assessment_id, assessment_name):
    @staticmethod
    def exam_to_pdf(exam, name, username):
        """
        Convert an exam assessment in JSON format to a PDF document.

        Parameters:
        - exam (dict): The exam assessment in dictionary format.

        Note:
        - This method creates a PDF document with formatted content based on the exam assessment.
        - The PDF is saved to the 'Files' directory with the naming convention 'exam.pdf'.
        - The method supports different types of exam sections, each with specific formatting for questions.
        - The exam dictionary should have the structure consistent with the expected format.
        - If the vertical position (y_position) exceeds the y_limit, a new page is added.
        """

        # Create a PDF document
        pdf_canvas = canvas.Canvas(rf"backend\api\media\{username}\exports\{name}_exam.pdf", pagesize=letter)

        # Add a header to the PDF
        pdf_canvas.setFont("Helvetica-Bold", 14)
        pdf_canvas.drawString(50, 770, "Exam")
        pdf_canvas.setFont("Helvetica", 12)

        y_position = 730
        x_position = 50
        line_height = 10
        max_line_length = 80
        y_limit = 100  # Set the y_limit

        for section in exam["sections"]:
            # Add section name to the PDF
            pdf_canvas.setFont("Helvetica-Bold", 12)
            pdf_canvas.drawString(x_position, y_position, section["section_name"])
            y_position -= line_height + 10  # Adjust the vertical position for the section name
            pdf_canvas.drawString(x_position, y_position, section["section_type"])
            pdf_canvas.setFont("Helvetica", 12)
            y_position -= line_height + 10  # Adjust the vertical position for the section name
            pdf_canvas.drawString(x_position, y_position, section["section_description"])
            y_position -= line_height # Adjust the vertical position for the section name

            questions = section["questions"]
            for index, question in enumerate(questions, start=1):
                y_position -= line_height  # Adjust the vertical position for each question

                question_text = question.get("question", "")
                wrapped_question_lines = Converter.wrap_text(f"{index}. {question_text}", max_line_length)
                for i, wrapped_question_line in enumerate(wrapped_question_lines):
                    if i == 0:
                        y_position -= line_height
                        pdf_canvas.drawString(x_position, y_position, wrapped_question_line)
                    else:
                        y_position -= 5
                        pdf_canvas.drawString(x_position + 15, y_position, wrapped_question_line)
                    y_position -= line_height

                # Add options to the PDF if it's a multiple choice section
                if section["section_type"] == "multiple choice":
                    options = question.get("options", [])
                    for option_index, option in enumerate(options, start=1):
                        y_position -= line_height  # Adjust the vertical position for each option
                        wrapped_option_lines = Converter.wrap_text(f"{chr(96 + option_index)}. {option}", max_line_length)
                        for i, wrapped_option_line in enumerate(wrapped_option_lines):
                            if i == 0:
                                y_position -= line_height
                            pdf_canvas.drawString(x_position + 15, y_position, wrapped_option_line)
                            y_position -= line_height

                y_position -= line_height # Adjust the vertical position for the next question

                # Check if y_position exceeds the y_limit, add a new page if needed
                if y_position < y_limit:
                    pdf_canvas.showPage()
                    y_position = 750  # Reset y_position for the new page

            y_position -= line_height + 10 # Adjust the vertical position for the next section

        # Save the PDF
        pdf_canvas.save()

    # If ever custom file naming convention is necessary
    # added username assessment id, and assessment_name for the file name to avoid errors
    # file naming convention: {assessment_name}_exam-answer-key_{username}-{assessment_id}.pdf
    # def exam_answer_key(exam, username, assessment_id, assessment_name):
    @staticmethod
    def exam_answer_key(exam, name, username):

        """
        Generate an answer key PDF for an exam assessment.

        Parameters:
        - exam (dict): The exam assessment in dictionary format.

        Note:
        - This method creates an answer key PDF document with correct answers for the exam assessment.
        - The PDF is saved to the 'Project Files' directory with the naming convention '{name}_exam-answer-key.pdf'.
        - The method supports different types of exam sections, each with specific formatting for correct answers.
        - The exam dictionary should have the structure consistent with the expected format.
        - If the vertical position (y_position) exceeds the y_limit, a new page is added.
        """
        #           ! SUGGESTIONS !
        # Can try to add section name and description ?
        # exam['name'] and exam['description'] 

        # Create a PDF document for the answer key
        pdf_canvas = canvas.Canvas(rf"backend\api\media\{username}\exports\{name}_exam_answer-key.pdf", pagesize=letter)

        # Add content to the PDF
        pdf_canvas.setFont("Helvetica", 12)
        max_line_length = 80  # Adjusted maximum line length
        y_position = 760
        y_limit = 100  # Set the y_limit
        line_height = 10  # Set the line height
        x_position = 80  # Adjusted starting position on the x-axis

        for section in exam["sections"]:
            # Add section name to the PDF
            pdf_canvas.setFont("Helvetica-Bold", 14)
            if section["section_type"] != "essay":
                pdf_canvas.drawString(50, y_position, section["section_type"])
            pdf_canvas.setFont("Helvetica", 12)
            y_position -= line_height  # Adjust the vertical position for the section name

            questions = section["questions"]
            for index, question in enumerate(questions, start=1):
                y_position -= 2 * line_height  # Adjust the vertical position for each question

                # Check if y_position exceeds the y_limit, add a new page if needed
                if y_position < y_limit:
                    pdf_canvas.showPage()
                    y_position = 780  # Reset y_position for the new page

                if section["section_type"] == "multiple choice":
                    correct_answer = f"Question {index}: {question['answer']}"
                elif section["section_type"] == "identification":
                    correct_answer = f"Question {index}: {question['answer']}"
                elif section["section_type"] == "true or false":
                    correct_answer = f"Question {index}: {'True' if question['answer'] else 'False'}"
                elif section["section_type"] == "fill in the blanks":
                    correct_answer = f"Question {index}: {question['answer']}"
                elif section["section_type"] == "essay":
                    pass
                
                if(section["section_type"] != "essay"):
                    wrapped_correct_answer = Converter.wrap_text(correct_answer, max_line_length)
                    for i, wrapped_question_line in enumerate(wrapped_correct_answer):
                        if i == 0:
                            y_position -= line_height
                            pdf_canvas.drawString(x_position, y_position, wrapped_question_line)
                        else:
                            y_position -= 5
                            pdf_canvas.drawString(x_position + 65, y_position, wrapped_question_line)
                        y_position -= line_height

            y_position -= 2 * line_height  # Adjust the vertical position for the next section

        # Save the PDF
        pdf_canvas.save()

    @staticmethod
    def quiz_to_gift(json_data, name, username):
        """
        Convert a quiz assessment in JSON format to a GIFT (General Import Format Template) file.

        Parameters:
        - json_data (dict): The quiz assessment in dictionary format.
        - output_file (str): The path to the output GIFT file.

        Note:
        - This method supports various types of quiz questions (multiple choice, identification, true or false, fill in the blanks, essay).
        - The GIFT file is created based on the input JSON data and saved to the specified output file.
        """

        gift_string = ""

        for question_data in json_data["questions"]:
            question_type = json_data["type"]
            question_text = question_data["question"]

            if question_type == "multiple choice":
                options = question_data.get("options", [])
                correct_answer_index = question_data.get("answer", 0)

                gift_string += f"::Question::{question_text}?\n"
                for i, option in enumerate(options):
                    if i == correct_answer_index:
                        gift_string += f"= {option}\n"
                    else:
                        gift_string += f"~ {option}\n"

            elif question_type == "identification":
                correct_answer = question_data.get("answer", "")

                gift_string += f"::Question::{question_text}?\n= {correct_answer}\n"

            elif question_type == "true or false":
                correct_answer = question_data.get("answer", False)

                gift_string += f"::Question::{question_text}?\n"
                if correct_answer:
                    gift_string += "= True\n~ False\n"
                else:
                    gift_string += "= False\n~ True\n"

            elif question_type == "fill in the blanks":
                correct_answer = question_data.get("answer", "")

                gift_string += f"::Question::{question_text} is ___?\n= {correct_answer}\n"

            elif question_type == "essay":
                gift_string += f"::Question::{question_text}?\n"

        # Save the GIFT content to the specified output file
        with open(rf"backend\api\media\{username}\exports\{name}_quiz-gift.txt", "w") as file:
            file.write(gift_string)

    @staticmethod
    def exam_to_gift(exam, name, username):
        """
        Convert an exam assessment in JSON format to a GIFT (General Import Format Template) file.
        
        Parameters:
        - exam (dict): The exam assessment in dictionary format.

        Note:
        - This method supports various types of exam sections (multiple choice, identification, true or false, fill in the blanks, essay).  
        - The GIFT file is created based on the input JSON data and saved to the specified output file.
        """
        
        gift_string = ""

        for section in exam["sections"]:
            section_name = section["section_name"]
            section_type = section["section_type"]
            questions = section["questions"]

            gift_string += f"::Section::{section_name}::{section_type}::\n"

            for question_data in questions:
                question_text = question_data["question"]

                if section_type == "multiple choice":
                    options = question_data.get("options", [])
                    correct_answer_index = question_data.get("answer", 0)

                    gift_string += f"::Question::{question_text}?\n"
                    for i, option in enumerate(options):
                        if i == correct_answer_index:
                            gift_string += f"= {option}\n"
                        else:
                            gift_string += f"~ {option}\n"

                elif section_type == "identification":
                    correct_answer = question_data.get("answer", "")

                    gift_string += f"::Question::{question_text}?\n= {correct_answer}\n"

                elif section_type == "true or false":
                    correct_answer = question_data.get("answer", False)

                    gift_string += f"::Question::{question_text}?\n"
                    if correct_answer:
                        gift_string += "= True\n~ False\n"
                    else:
                        gift_string += "= False\n~ True\n"

                elif section_type == "fill in the blanks":
                    correct_answer = question_data.get("answer", "")

                    gift_string += f"::Question::{question_text} is ___?\n= {correct_answer}\n"

                elif section_type == "essay":
                    gift_string += f"::Question::{question_text}?\n"

        # Save the GIFT content to the specified output file
        with open(rf"backend\api\media\{username}\exports\{name}_exam-gift.txt", "w") as file:
            file.write(gift_string)

        return gift_string
    
    @staticmethod
    def quiz_to_docx(quiz, name, username):
        document = Document()
        
        # Get the type of the quiz
        quiz_type = quiz["type"]
        document.add_heading(quiz_type, level=1)

        match(quiz_type.lower()):
            case "multiple choice":
                # Add questions to the document
                for i, question_data in enumerate(quiz["questions"], start=1):
                    question_text = f"{i}. {question_data['question']}"
                    document.add_paragraph(question_text, style="Body Text")

                    for j, option in enumerate(question_data["options"], start=1):
                        option_text = f"  {chr(64 + j)}. {option}"
                        document.add_paragraph(option_text)

                    # Add space between questions
                    document.add_paragraph()

                # Add page break before answers
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

                # Add answers
                document.add_heading("Answer Key", level=1)
                for i, question_data in enumerate(quiz["questions"], start=1):
                    answer_text = f"Question {i}: {question_data['answer']}"
                    document.add_paragraph(answer_text, style="Body Text")
            
            case "essay":
                for i, question_data in enumerate(quiz["questions"], start=1):
                    question_text = f"{i}. {question_data['question']}"
                    document.add_paragraph(question_text, style="Body Text")

            case "identification":
                for i, question_data in enumerate(quiz["questions"], start=1):
                    question_text = f"________ {i}. {question_data['question']}"
                    document.add_paragraph(question_text, style="Body Text")

                # Add page break before answers
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

                # Add answers
                document.add_heading("Answer Key", level=1)
                for i, question_data in enumerate(quiz["questions"], start=1):
                    correct_answer = question_data["answer"]
                    answer_text = f"Question {i}: {correct_answer}"
                    document.add_paragraph(answer_text, style="Body Text")
        
            case "true or false" | "fill in the blanks":
                for i, question_data in enumerate(quiz["questions"], start=1):
                    question_text = f"{i}. {question_data['question']}"
                    document.add_paragraph(question_text, style="Body Text")

                # Add page break before answers
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

                # Add answers
                document.add_heading("Answer Key", level=1)
                for i, question_data in enumerate(quiz["questions"], start=1):
                    correct_answer = question_data["answer"]
                    answer_text = f"Question {i}: {correct_answer}"
                    document.add_paragraph(answer_text, style="Body Text")

        # Save the document
        file_path = fr'backend\api\media\{username}\exports\{name}_quiz.docx'
        document.save(file_path)

    @staticmethod
    def exam_to_docx(exam, name, username):

        document = Document()

        document.add_heading(exam["type"], level=1)

        for section in exam["sections"]:
            print(section)
            section_name = section.get("section_name", "") or section.get("name", "")
            section_type = section.get("section_type", "") or section.get("type", "")
            document.add_heading(f"{section_name} - {section_type}", level=2)

            for i, question_data in enumerate(section["questions"], start=1):
                question_text = f"{i}. {question_data['question']}"
                document.add_paragraph(question_text, style="Heading 2")

                if section_type.lower() == "multiple choice":
                    print("multiple choice Questions \n\n")
                    for j, option in enumerate(question_data["options"], start=1):
                        option_text = f"  {chr(64 + j)}. {option}"
                        document.add_paragraph(option_text)
                # Add space between questions
                document.add_paragraph()
            
        # Add page break before answers
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Add answers
        document.add_heading("Answer Key", level=1)
        for section in exam["sections"]:
            section_name = section.get("name", "") or section.get("section_name", "")
            section_type = section.get("type", "") or section.get("section_type", "")
            
            if section_type.lower() != "essay":    
                
                document.add_heading(f"{section_name} - {section_type}", level=2)

                if section_type.lower() == "multiple choice":
                    for i, answer_data in enumerate(section["questions"], start=1):
                        correct_answer = answer_data["answer"]
                        answer_text = f"{i}: {correct_answer}"
                        document.add_paragraph(answer_text, style="Body Text")
                else:
                    for i, answer_data in enumerate(section["questions"], start=1):
                        question_text = f"{i}. {answer_data['answer']}"
                        document.add_paragraph(question_text, style="Body Text")

        print('Here')
        # Save the document
        file_path = fr'backend\api\media\{username}\exports\{name}_exam.docx'
        document.save(file_path)